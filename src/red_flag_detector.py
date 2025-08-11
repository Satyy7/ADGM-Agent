
import json
import shutil
from pathlib import Path
from typing import Optional, List

import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

from src.utils import gemini_generate

PROMPT_TEMPLATE = """
Act as an ADGM Compliance Officer. Analyze the 'Uploaded Document Text' against the 'Reference Text' for legal red flags.

**Format your output as a JSON object with a single key "issues_found" which is a list of issue objects.** Each object MUST have the following keys:
- "section": The clause number or general section title where the issue was found (e.g., "Appointment of Director(s)").
- "problematic_text": The exact, verbatim text snippet from the document that contains the issue (e.g., "disputes shall be resolved in UAE Federal Courts").
- "issue": A concise description of the red flag (e.g., "Incorrect Jurisdiction").
- "citation": The specific ADGM law or rule that applies (e.g., "Per ADGM Companies Regulations 2020, Art. 17..."). If no specific article applies, cite the general principle.
- "severity": "High", "Medium", or "Low".
- "suggestion": A concrete recommendation on how to fix the issue.

**Do not include any text outside the JSON object.**

---
**Reference Text (Official ADGM Regulations/Templates):**
{reference_text}
---
**Uploaded Document Text:**
{uploaded_text}
---

**JSON Output:**
"""

def detect_red_flags(uploaded_text: str, reference_text: str) -> list:
    """
    Calls your LLM (gemini_generate) with the prompt and returns the 'issues_found' list.
    (Unchanged behavior from your original code.)
    """
    prompt = PROMPT_TEMPLATE.format(
        reference_text=reference_text or "No official reference text was found for comparison.",
        uploaded_text=uploaded_text[:15000]
    )
    try:
        response_text = gemini_generate(prompt)
        clean_json_str = response_text.strip().replace("```json", "").replace("```", "")
        result = json.loads(clean_json_str)
        return result.get("issues_found", [])
    except (json.JSONDecodeError, Exception) as e:
        print(f"[red_flag_detector] Error processing LLM response for red flag detection: {e}")
        return []



def _normalize_whitespace(s: str) -> str:
    return " ".join(s.split()) if s is not None else ""

def _find_paragraph_by_text(doc: docx.document.Document, problematic_text: Optional[str], section_text: Optional[str]) -> Optional[docx.text.paragraph.Paragraph]:
    """
    Try multiple strategies to find a paragraph:
      1. exact substring match
      2. case-insensitive match
      3. normalized whitespace match
      4. token overlap (simple fallback)
      5. section_text matching as last resort
    Returns the paragraph object or None.
    """
    if problematic_text:
        p_text = problematic_text.strip()
        if not p_text:
            problematic_text = None

    
    if problematic_text:
        lower_target = problematic_text.lower()
        for paragraph in doc.paragraphs:
            try:
                if problematic_text in paragraph.text:
                    return paragraph
            except Exception:
                pass
        for paragraph in doc.paragraphs:
            try:
                if lower_target in paragraph.text.lower():
                    return paragraph
            except Exception:
                pass

   
    if problematic_text:
        norm_target = _normalize_whitespace(problematic_text).lower()
        for paragraph in doc.paragraphs:
            if _normalize_whitespace(paragraph.text).lower().find(norm_target) != -1:
                return paragraph

    
    if problematic_text:
        import re
        tokens = re.findall(r"\w+", problematic_text.lower())
        if tokens:
            token_set = set(tokens)
            for paragraph in doc.paragraphs:
                p_tokens = set(re.findall(r"\w+", paragraph.text.lower()))
                if not p_tokens:
                    continue
                
                inter = token_set.intersection(p_tokens)
                if len(inter) / max(1, len(token_set)) >= 0.6:
                    return paragraph

    
    if section_text:
        st = section_text.strip()
        for paragraph in doc.paragraphs:
            if st in paragraph.text:
                return paragraph
        for paragraph in doc.paragraphs:
            if st.lower() in paragraph.text.lower():
                return paragraph
        st_norm = _normalize_whitespace(st).lower()
        for paragraph in doc.paragraphs:
            if _normalize_whitespace(paragraph.text).lower().find(st_norm) != -1:
                return paragraph

    return None



def add_comments_to_docx(original_docx_path: Path, issues: List[dict], output_docx_path: Path, debug: bool = False):
    """
    For each issue:
      - try to find the paragraph using multiple heuristics
      - if problematic_text is found inside the paragraph, reconstruct the paragraph
        as: [before][highlighted_problematic_text][after]
      - insert an annotation paragraph directly AFTER the flagged paragraph with:
        "[COMMENT by Corporate Agent] [SEVERITY] issue | Legal Basis: citation | Suggestion: suggestion"
      - if matching fails, try to add the annotation at the end of document (so nothing is lost)
    """
    try:
        doc = docx.Document(str(original_docx_path))

        for idx, issue in enumerate(issues):
            problem_text = issue.get('problematic_text')
            section_text = issue.get('section')
            if debug:
                print(f"[add_comments] processing issue {idx}: section='{section_text}' problem_text='{problem_text}'")

            target_paragraph = _find_paragraph_by_text(doc, problem_text, section_text)

            if target_paragraph is None:
                
                comment_text = (
                    f"[COMMENT by Corporate Agent] [{issue.get('severity','Unknown').upper()}] {issue.get('issue','No description.')} "
                    f"| Legal Basis: {issue.get('citation','No citation.')} | Suggestion: {issue.get('suggestion','No suggestion.')}"
                )
                if debug:
                    print(f"[add_comments] no match found - appending comment to document end: {comment_text}")
                new_p = doc.add_paragraph(comment_text)
                for run in new_p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                continue

           
            para_text = target_paragraph.text or ""
            matched_and_highlighted = False

            if problem_text:
                
                low_para = para_text.lower()
                low_problem = problem_text.lower()
                pos = low_para.find(low_problem)
                if pos != -1:
                   
                    start = pos
                    end = pos + len(problem_text)  
                   
                    before = para_text[:start]
                    mid = para_text[start:start + len(problem_text)]
                    after = para_text[start + len(problem_text):]
                   
                    target_paragraph.text = before  
                    run_mid = target_paragraph.add_run(mid)
                    try:
                        run_mid.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except Exception:
                      
                        pass
                    try:
                        run_mid.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    except Exception:
                        pass
                   
                    if after:
                        target_paragraph.add_run(after)
                    matched_and_highlighted = True
                    if debug:
                        print(f"[add_comments] highlighted problem_text in paragraph: '{mid}'")
                else:
                    
                    norm_para = _normalize_whitespace(para_text).lower()
                    norm_problem = _normalize_whitespace(problem_text).lower()
                    pos2 = norm_para.find(norm_problem)
                    if pos2 != -1:
                      
                        tokens = problem_text.split()
                        
                        first_tok = tokens[0].lower()
                        tok_pos = low_para.find(first_tok)
                        if tok_pos != -1:
                            
                            approx_len = sum(len(t) for t in tokens) + max(0, len(tokens)-1)  
                            before = para_text[:tok_pos]
                            mid = para_text[tok_pos:tok_pos+approx_len]
                            after = para_text[tok_pos+approx_len:]
                            target_paragraph.text = before
                            run_mid = target_paragraph.add_run(mid)
                            try:
                                run_mid.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            except Exception:
                                pass
                            try:
                                run_mid.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                            except Exception:
                                pass
                            if after:
                                target_paragraph.add_run(after)
                            matched_and_highlighted = True
                            if debug:
                                print(f"[add_comments] approximate highlight applied: '{mid}'")

            
            comment_text = (
                f"[COMMENT by Corporate Agent] [{issue.get('severity','Unknown').upper()}] {issue.get('issue','No description.')} "
                f"| Legal Basis: {issue.get('citation','No citation.')} | Suggestion: {issue.get('suggestion','No suggestion.')}"
            )

            
            try:
                new_para = doc.add_paragraph(comment_text)
                
                target_paragraph._p.addnext(new_para._p)
               
                for run in new_para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            except Exception as e:
                
                if debug:
                    print(f"[add_comments] could not insert after paragraph due to: {e}. Appending at end instead.")
                new_para = doc.add_paragraph(comment_text)
                for run in new_para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        
        doc.save(str(output_docx_path))

    except Exception as e:
        print(f"[add_comments_to_docx] Critical error while writing DOCX: {e}")
        
        try:
            shutil.copy(str(original_docx_path), str(output_docx_path))
        except Exception as e2:
            print(f"[add_comments_to_docx] Also failed to copy original: {e2}")
